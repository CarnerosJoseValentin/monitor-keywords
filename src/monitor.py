import os
import re
import json
import time
import hashlib
import unicodedata
from datetime import datetime, timedelta, timezone
from urllib.parse import urlparse, parse_qsl, urlencode, urlunparse
import feedparser
import requests
from dateutil import parser as dateparser

# Nuevas librerías para el reporte
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

STATE_PATH = "state.json"
KEYWORDS_PATH = "keywords.txt"
FEEDS_PATH = "feeds.txt"

DEFAULT_RETENTION_DAYS = 30
DEFAULT_MAX_SNIPPET = 300

TRACKING_PARAMS_PREFIX = ("utm_",)
TRACKING_PARAMS_EXACT = {
    "fbclid", "gclid", "igshid", "mc_cid", "mc_eid", "ref", "ref_src",
    "mkt_tok", "spm", "yclid"
}

UA = "keyword-monitor-bot/1.0 (+GitHub Actions)"


def read_lines(path: str) -> list[str]:
    if not os.path.exists(path):
        return []
    out = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            s = line.strip()
            if not s or s.startswith("#"):
                continue
            out.append(s)
    return out


def strip_quotes(s: str) -> str:
    s = s.strip()
    if (s.startswith('"') and s.endswith('"')) or (s.startswith("'") and s.endswith("'")):
        return s[1:-1].strip()
    return s


def fold_accents(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    return s


def normalize_text_for_match(s: str) -> str:
    s = strip_quotes(s)
    s = s.lower()
    s = fold_accents(s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def normalize_url(url: str) -> str:
    try:
        p = urlparse(url.strip())
        fragment = ""
        q = []
        for k, v in parse_qsl(p.query, keep_blank_values=True):
            kl = k.lower()
            if kl in TRACKING_PARAMS_EXACT:
                continue
            if any(kl.startswith(pref) for pref in TRACKING_PARAMS_PREFIX):
                continue
            q.append((k, v))
        query = urlencode(q, doseq=True)
        netloc = p.netloc.lower()
        path = p.path or "/"
        if path != "/" and path.endswith("/"):
            path = path[:-1]
        return urlunparse((p.scheme, netloc, path, p.params, query, fragment))
    except Exception:
        return url.strip()


def stable_id(title: str, url: str, published: str | None) -> str:
    base = (title or "").strip().lower() + "||" + normalize_url(url).lower()
    if published:
        base += "||" + published.strip()
    return hashlib.sha256(base.encode("utf-8")).hexdigest()


def load_state() -> dict:
    if not os.path.exists(STATE_PATH):
        return {"version": 1, "seen": {}, "reporte_diario": []}
    with open(STATE_PATH, "r", encoding="utf-8") as f:
        state = json.load(f)
        if "reporte_diario" not in state:
            state["reporte_diario"] = []
        return state


def save_state(state: dict) -> None:
    with open(STATE_PATH, "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2, sort_keys=True)


def purge_state(state: dict, retention_days: int) -> None:
    cutoff = datetime.now(timezone.utc) - timedelta(days=retention_days)
    seen = state.get("seen", {})
    to_del = []
    for k, v in seen.items():
        try:
            dt = dateparser.parse(v)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            if dt < cutoff:
                to_del.append(k)
        except Exception:
            to_del.append(k)
    for k in to_del:
        del seen[k]
    state["seen"] = seen


def entry_text(entry) -> str:
    parts = []
    for key in ("title", "summary", "description"):
        if key in entry and entry[key]:
            parts.append(str(entry[key]))
    if "content" in entry:
        try:
            for c in entry["content"]:
                if "value" in c and c["value"]:
                    parts.append(str(c["value"]))
        except Exception:
            pass
    text = "\n".join(parts)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def compile_keyword_patterns(keywords: list[str]) -> list[tuple[str, re.Pattern]]:
    compiled = []
    for raw in keywords:
        kw = strip_quotes(raw)
        kw_norm = normalize_text_for_match(kw)
        if not kw_norm:
            continue
        is_token = bool(re.fullmatch(r"[a-z0-9_]+", kw_norm))
        if is_token and len(kw_norm) <= 6:
            pat = re.compile(rf"\b{re.escape(kw_norm)}\b", re.IGNORECASE)
        else:
            pat = re.compile(re.escape(kw_norm), re.IGNORECASE)
        compiled.append((kw, pat))
    return compiled


def match_keywords(text: str, compiled: list[tuple[str, re.Pattern]]) -> list[str]:
    t = normalize_text_for_match(text)
    hits = []
    for kw_original, pat in compiled:
        if pat.search(t):
            hits.append(kw_original)
    return hits


def parse_published(entry) -> str | None:
    for key in ("published", "updated"):
        if key in entry and entry[key]:
            return str(entry[key])
    return None


def fetch_feed(url: str) -> feedparser.FeedParserDict:
    return feedparser.parse(url, agent=UA)


def send_telegram(token: str, chat_id: str, text: str) -> None:
    url = f"https://api.telegram.org/bot{token}/sendMessage"
    payload = {
        "chat_id": chat_id,
        "text": text,
        "disable_web_page_preview": False,
    }
    r = requests.post(url, json=payload, timeout=30)
    if not r.ok:
        print(f"Telegram sendMessage failed: {r.status_code} {r.text}")


def send_telegram_document(token: str, chat_id: str, file_path: str, caption: str = "") -> None:
    url = f"https://api.telegram.org/bot{token}/sendDocument"
    with open(file_path, "rb") as f:
        files = {"document": f}
        data = {"chat_id": chat_id, "caption": caption}
        r = requests.post(url, data=data, files=files, timeout=60)
        if not r.ok:
            print(f"Telegram sendDocument failed: {r.status_code} {r.text}")


def generar_docx(items: list, api_key: str, path="reporte_diario.docx"):
    doc = Document()
    now_arg = datetime.now(timezone.utc) - timedelta(hours=3)
    meses = ["ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO", "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"]
    fecha_encabezado = f"{now_arg.day} {meses[now_arg.month - 1]} {now_arg.year}"

    # Encabezado
    titulo = doc.add_heading("REPORTE DE EXPLOTACIÓN DE PRENSA – (REP)", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.add_run(fecha_encabezado).bold = True
    doc.add_paragraph()

    for idx, item in enumerate(items):
        if idx > 0:
            doc.add_paragraph("-" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # Extraer y formatear datos
        dt_pub = now_arg 
        if item.get("published"):
            try:
                dt_pub = dateparser.parse(item["published"])
                if dt_pub.tzinfo:
                    dt_pub = dt_pub.astimezone(timezone(timedelta(hours=-3)))
            except:
                pass

        fecha_hecho = dt_pub.strftime("%d/%m/%Y")
        hora_hecho = dt_pub.strftime("%H:%M hs")
        delitos = ", ".join(item.get("keywords", [])).title()
        titulo_noticia = item.get("title", "")
        url = item.get("url", "")
        texto_base = item.get("text", "")
        
        # Detectar provincia en el texto
        prov = "Provincia a determinar"
        txt_lower = texto_base.lower()
        for pr in ["Córdoba", "Mendoza", "San Luis", "San Juan", "La Pampa"]:
            if pr.lower() in txt_lower:
                prov = pr
                break

        # Procesar con Gemini
        resumen = texto_base[:500] + "..." # Fallback por si falla la API
        if api_key:
            try:
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('gemini-1.5-flash')
                prompt = f"Redacta un resumen formal, objetivo y con tono institucional policial de la siguiente noticia. Debe ser un solo párrafo estructurado, directo al punto:\n\n{texto_base}"
                response = model.generate_content(prompt)
                resumen = response.text.strip()
            except Exception as e:
                print(f"Error procesando con Gemini: {e}")

        # Escribir en el documento
        p = doc.add_paragraph()
        p.add_run("Ámbito: ").bold = True
        p.add_run("URSA II DEL CENTRO\n")

        p.add_run("Fecha del hecho: ").bold = True
        p.add_run(f"{fecha_hecho}\n")

        p.add_run("Hora: ").bold = True
        p.add_run(f"{hora_hecho}\n")

        p.add_run("Provincia: ").bold = True
        p.add_run(f"{prov}\n")

        p.add_run("Delito: ").bold = True
        p.add_run(f"{delitos}\n")

        p.add_run("Título: ").bold = True
        p.add_run(f"{titulo_noticia}\n")

        p.add_run("Resumen:\n").bold = True
        p.add_run(f"{resumen}\n\n")

        p.add_run("Fuente: ").bold = True
        p.add_run(url)

    doc.save(path)
    return path


def main():
    token = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()
    chat_id = os.getenv("TELEGRAM_CHAT_ID", "").strip()
    api_key_gemini = os.getenv("GEMINI_API_KEY", "").strip()

    if not token or not chat_id:
        raise SystemExit("Faltan credenciales de Telegram.")

    retention_days = int(os.getenv("RETENTION_DAYS", str(DEFAULT_RETENTION_DAYS)))
    max_snippet = int(os.getenv("MAX_SNIPPET_CHARS", str(DEFAULT_MAX_SNIPPET)))

    keywords_raw = read_lines(KEYWORDS_PATH)
    compiled = compile_keyword_patterns(keywords_raw)
    all_feeds = read_lines(FEEDS_PATH)

    state = load_state()
    purge_state(state, retention_days)
    seen = state.get("seen", {})
    reporte_diario = state.get("reporte_diario", [])

    new_items = []

    for feed_url in all_feeds:
        try:
            parsed = fetch_feed(feed_url)
            for entry in parsed.entries:
                title = getattr(entry, "title", "") or entry.get("title", "") or ""
                link = getattr(entry, "link", "") or entry.get("link", "") or ""
                if not link:
                    continue
                norm = normalize_url(link)
                text = (title + " " + entry_text(entry)).strip()
                hits = match_keywords(text, compiled)
                if not hits:
                    continue

                published = parse_published(entry)
                sid = stable_id(title, norm, published)

                if sid in seen:
                    continue

                snippet_src = entry_text(entry)
                snippet = snippet_src[:max_snippet] + "…" if len(snippet_src) > max_snippet else snippet_src
                domain = urlparse(norm).netloc

                item_data = {
                    "keywords": hits,
                    "title": title.strip() or "(sin título)",
                    "url": norm,
                    "domain": domain,
                    "snippet": snippet,
                    "text": text,
                    "published": published
                }
                new_items.append(item_data)
                reporte_diario.append(item_data)
                seen[sid] = datetime.now(timezone.utc).isoformat()

        except Exception as e:
            print(f"Error en feed {feed_url}: {e}")

    # Enviar Alertas Horarias Estándar
    new_items_sorted = sorted(new_items, key=lambda it: (it["domain"], it["title"].lower()))
    if new_items_sorted:
        lines = [f"🔔 Nuevos hallazgos (última hora): {len(new_items_sorted)}\n"]
        for it in new_items_sorted[:50]:
            kws = ", ".join(it["keywords"][:5])
            lines.append(f"• 🧷 {kws}")
            lines.append(f"  📰 {it['title']}")
            lines.append(f"  🌐 {it['domain']}")
            lines.append(f"  🔗 {it['url']}\n")

        msg = "\n".join(lines).strip()
        chunks = [msg[i:i+3900] for i in range(0, len(msg), 3900)]
        for chunk in chunks:
            send_telegram(token, chat_id, chunk)
            time.sleep(1)

    # Actualizar estado de acumulación
    state["seen"] = seen
    state["reporte_diario"] = reporte_diario

    # Lógica de las 10:00 AM (Hora Argentina UTC-3)
    now_arg = datetime.now(timezone.utc) - timedelta(hours=3)
    today_str = now_arg.strftime("%Y-%m-%d")

    # Si es la ejecución de las 10 AM y el reporte de hoy aún no se envió
    if now_arg.hour == 9 and state.get("last_report_date") != today_str:
        if len(reporte_diario) > 0:
            try:
                docx_path = generar_docx(reporte_diario, api_key_gemini)
                caption = f"📄 REPORTE DE EXPLOTACIÓN DE PRENSA (REP)\nCorrespondiente al periodo finalizado el {fecha_hecho if 'fecha_hecho' in locals() else today_str} a las 10:00 hs."
                send_telegram_document(token, chat_id, docx_path, caption)
                
                # Vaciar la caja fuerte para las siguientes 24 horas y marcar como enviado
                state["reporte_diario"] = []
                state["last_report_date"] = today_str
            except Exception as e:
                print(f"Error al enviar el Reporte Diario: {e}")
        else:
            # Marcar igual para no re-intentar todo el día si no hubo noticias
            state["last_report_date"] = today_str 

    save_state(state)


if __name__ == "__main__":
    main()
